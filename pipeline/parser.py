"""
D2S Pipeline — DocumentParser: extract elements from DOCX and PDF.
"""
import io
from dataclasses import dataclass, field
from enum import Enum
from typing import Optional

from logger import logger


class ElementType(Enum):
    HEADING = "heading"
    PARAGRAPH = "paragraph"
    TABLE = "table"
    IMAGE = "image"


@dataclass
class DocumentElement:
    type: ElementType
    content: str = ""
    heading_level: int = 0  # 1, 2, 3 for headings
    table_data: list[list[str]] = field(default_factory=list)
    image_bytes: Optional[bytes] = None
    order: int = 0


def parse_docx(file_bytes: bytes) -> list[DocumentElement]:
    """Parse DOCX file using python-docx, iterating body in document order."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(file_bytes))
    elements: list[DocumentElement] = []
    order = 0

    # Build lookup maps keyed by XML element identity for O(1) access
    para_idx = {id(p._element): p for p in doc.paragraphs}
    table_idx = {id(t._element): t for t in doc.tables}

    # Iterate body children in actual document order
    for child in doc.element.body.iterchildren():
        tag = child.tag

        if tag == qn('w:p'):
            # ── Paragraph (may contain text and/or inline images) ──
            para = para_idx.get(id(child))
            if not para:
                continue

            text = para.text.strip()
            style_name = (para.style.name or "").lower()

            if text:
                if "heading" in style_name:
                    level = 1
                    for i in range(1, 4):
                        if str(i) in style_name:
                            level = i
                            break
                    elements.append(DocumentElement(
                        type=ElementType.HEADING,
                        content=text,
                        heading_level=level,
                        order=order,
                    ))
                else:
                    elements.append(DocumentElement(
                        type=ElementType.PARAGRAPH,
                        content=text,
                        order=order,
                    ))
                order += 1

            # Check for inline images (a:blip) inside this paragraph
            for blip in child.findall('.//' + qn('a:blip')):
                rId = blip.get(qn('r:embed'))
                if rId and rId in doc.part.rels:
                    rel = doc.part.rels[rId]
                    if "image" in rel.reltype:
                        try:
                            image_data = rel.target_part.blob
                            elements.append(DocumentElement(
                                type=ElementType.IMAGE,
                                image_bytes=image_data,
                                order=order,
                            ))
                            order += 1
                        except Exception as e:
                            logger.warning("Failed to extract DOCX image: %s", e)

        elif tag == qn('w:tbl'):
            # ── Table ──
            table = table_idx.get(id(child))
            if not table:
                continue

            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                elements.append(DocumentElement(
                    type=ElementType.TABLE,
                    table_data=rows,
                    order=order,
                ))
                order += 1

    logger.info("DOCX parsed: %d elements", len(elements))
    return elements


def parse_pdf(file_bytes: bytes) -> list[DocumentElement]:
    """Parse PDF file using PyMuPDF with font-size heuristic for headings.

    All element types (text, images, tables) are collected per page with their
    Y-position, then sorted so the final order matches the visual layout.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    elements: list[DocumentElement] = []
    order = 0
    page_count = len(doc)

    for page_num in range(page_count):
        page = doc[page_num]

        # Collect (y_position, element) tuples — sort by Y at end of page
        page_elements: list[tuple[float, DocumentElement]] = []

        # ── 1. Text blocks ──
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

        for block in blocks:
            if block["type"] != 0:  # only text blocks here
                continue
            for line in block["lines"]:
                text_parts = []
                max_font_size = 0
                for span in line["spans"]:
                    text_parts.append(span["text"])
                    max_font_size = max(max_font_size, span["size"])

                text = " ".join(text_parts).strip()
                if not text:
                    continue

                line_y = line["bbox"][1]

                # Font-size heuristic for headings
                if max_font_size > 16:
                    elem = DocumentElement(
                        type=ElementType.HEADING, content=text,
                        heading_level=1, order=0,
                    )
                elif max_font_size > 13:
                    elem = DocumentElement(
                        type=ElementType.HEADING, content=text,
                        heading_level=2, order=0,
                    )
                elif max_font_size > 11.5:
                    elem = DocumentElement(
                        type=ElementType.HEADING, content=text,
                        heading_level=3, order=0,
                    )
                else:
                    elem = DocumentElement(
                        type=ElementType.PARAGRAPH, content=text, order=0,
                    )

                page_elements.append((line_y, elem))

        # ── 2. Images via page.get_images() (correct xref-based API) ──
        seen_xrefs: set[int] = set()
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            if xref in seen_xrefs:
                continue
            seen_xrefs.add(xref)
            try:
                img_data = doc.extract_image(xref)
                if img_data and img_data.get("image"):
                    # Determine image Y-position on page
                    try:
                        img_rects = page.get_image_rects(xref)
                        img_y = img_rects[0].y0 if img_rects else 0.0
                    except Exception:
                        img_y = 0.0

                    elem = DocumentElement(
                        type=ElementType.IMAGE,
                        image_bytes=img_data["image"],
                        order=0,
                    )
                    page_elements.append((img_y, elem))
            except Exception as e:
                logger.warning(
                    "Failed to extract PDF image (xref=%d) on page %d: %s",
                    xref, page_num + 1, e,
                )

        # ── 3. Tables ──
        try:
            tables = page.find_tables()
            for table in tables:
                rows = []
                for row in table.extract():
                    cells = [str(cell) if cell is not None else "" for cell in row]
                    rows.append(cells)
                if rows:
                    table_y = table.bbox[1]  # top-Y of table bounding box
                    elem = DocumentElement(
                        type=ElementType.TABLE, table_data=rows, order=0,
                    )
                    page_elements.append((table_y, elem))
        except Exception as e:
            logger.warning("Failed to extract tables from page %d: %s", page_num + 1, e)

        # ── 4. Sort by Y-position and assign sequential order ──
        page_elements.sort(key=lambda x: x[0])
        for _, elem in page_elements:
            elem.order = order
            elements.append(elem)
            order += 1

    doc.close()
    logger.info("PDF parsed: %d elements from %d pages", len(elements), page_count)
    return elements


def parse_document(file_bytes: bytes, file_type: str) -> list[DocumentElement]:
    """Parse a document file and return ordered elements."""
    if file_type == "docx":
        return parse_docx(file_bytes)
    elif file_type == "pdf":
        return parse_pdf(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
